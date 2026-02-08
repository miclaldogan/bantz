"""Document Ingest v0 (Issue #460).

Unified pipeline for parsing, chunking, and querying documents.

Supported formats:
- **TXT** / **Markdown** — plain text
- **HTML** — stripped via basic tag removal
- **PDF** — via PyMuPDF (fitz) if available, else fallback
- **DOCX** — via python-docx if available, else fallback

Features
--------
- :class:`Document` dataclass with content, metadata, chunks
- Token-aware chunking with configurable overlap
- Sandbox enforcement: only allowed directories
- Symlink resolution for security
- ``summarize()`` and ``query()`` stubs (return heuristic summaries)

See Also
--------
- ``src/bantz/document/ingestion.py`` — legacy ingestion
- ``src/bantz/document/parsers/`` — format-specific parsers
"""

from __future__ import annotations

import html
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

logger = logging.getLogger(__name__)

__all__ = [
    "Document",
    "Chunk",
    "DocumentIngester",
    "IngestConfig",
]


# ── Dataclasses ───────────────────────────────────────────────────────

@dataclass
class Chunk:
    """A piece of a document."""

    index: int
    text: str
    start_char: int = 0
    end_char: int = 0
    estimated_tokens: int = 0


@dataclass
class Document:
    """Parsed document representation."""

    path: str
    format: str                 # "pdf" | "docx" | "txt" | "md" | "html"
    content: str                # Full text
    title: Optional[str] = None
    pages: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[Chunk] = field(default_factory=list)


# ── Config ────────────────────────────────────────────────────────────

@dataclass
class IngestConfig:
    """Configuration for the ingester."""

    chunk_size: int = 1000          # Max tokens per chunk
    chunk_overlap: int = 100        # Token overlap between chunks
    allowed_dirs: List[str] = field(default_factory=lambda: [
        str(Path.home() / "Documents"),
        str(Path.home() / "Downloads"),
        "/tmp",
    ])
    chars_per_token: float = 4.0    # Rough estimate


# ── Ingester ──────────────────────────────────────────────────────────

class DocumentIngester:
    """Parse, chunk, and query documents.

    Parameters
    ----------
    config:
        Ingest configuration.
    summarizer:
        Optional ``(text) → summary`` callback (e.g. LLM call).
        If ``None``, a simple truncation summarizer is used.
    """

    def __init__(
        self,
        config: Optional[IngestConfig] = None,
        summarizer: Optional[Callable[[str], str]] = None,
    ) -> None:
        self._config = config or IngestConfig()
        self._summarizer = summarizer or self._default_summarizer

    # ── ingest ────────────────────────────────────────────────────────

    def ingest(self, file_path: str) -> Document:
        """Parse a file and return a :class:`Document`.

        Raises
        ------
        PermissionError
            If file is outside allowed directories.
        FileNotFoundError
            If file does not exist.
        ValueError
            If format is not supported.
        """
        path = Path(file_path).expanduser()

        # Resolve symlinks for security
        resolved = path.resolve()
        if not resolved.exists():
            raise FileNotFoundError(f"File not found: {resolved}")

        self._check_sandbox(resolved)

        ext = resolved.suffix.lower().lstrip(".")
        fmt = self._normalize_format(ext)
        content = self._parse(resolved, fmt)
        title = resolved.stem

        doc = Document(
            path=str(resolved),
            format=fmt,
            content=content,
            title=title,
            metadata={"size_bytes": resolved.stat().st_size},
        )

        doc.chunks = self.chunk(doc)
        return doc

    # ── chunking ──────────────────────────────────────────────────────

    def chunk(self, doc: Document, chunk_size: Optional[int] = None) -> List[Chunk]:
        """Split document into token-aware chunks with overlap."""
        size = chunk_size or self._config.chunk_size
        overlap = self._config.chunk_overlap
        cpt = self._config.chars_per_token

        char_size = int(size * cpt)
        char_overlap = int(overlap * cpt)

        text = doc.content
        chunks: List[Chunk] = []
        start = 0
        idx = 0

        while start < len(text):
            end = min(start + char_size, len(text))
            chunk_text = text[start:end]

            chunks.append(Chunk(
                index=idx,
                text=chunk_text,
                start_char=start,
                end_char=end,
                estimated_tokens=int(len(chunk_text) / cpt),
            ))
            idx += 1
            start = end - char_overlap
            if start >= len(text):
                break
            if end == len(text):
                break

        return chunks

    # ── summarize ─────────────────────────────────────────────────────

    def summarize(self, doc: Document) -> str:
        """Return a summary of the document."""
        return self._summarizer(doc.content)

    # ── query ─────────────────────────────────────────────────────────

    def query(self, doc: Document, question: str) -> str:
        """Answer a question about the document (RAG-lite).

        Finds the most relevant chunk by simple keyword overlap,
        then summarizes it.
        """
        if not doc.chunks:
            return self._summarizer(doc.content)

        question_words = set(question.lower().split())
        best_chunk = max(
            doc.chunks,
            key=lambda c: len(question_words & set(c.text.lower().split())),
        )
        return self._summarizer(best_chunk.text)

    # ── parsing ───────────────────────────────────────────────────────

    def _parse(self, path: Path, fmt: str) -> str:
        """Dispatch to format-specific parser."""
        if fmt in ("txt", "md"):
            return self._parse_text(path)
        if fmt == "html":
            return self._parse_html(path)
        if fmt == "pdf":
            return self._parse_pdf(path)
        if fmt == "docx":
            return self._parse_docx(path)
        raise ValueError(f"Unsupported format: {fmt}")

    @staticmethod
    def _parse_text(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _parse_html(path: Path) -> str:
        raw = path.read_text(encoding="utf-8", errors="replace")
        # Strip tags
        text = re.sub(r"<[^>]+>", " ", raw)
        text = html.unescape(text)
        return re.sub(r"\s+", " ", text).strip()

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n".join(pages)
        except ImportError:
            # Fallback: read raw bytes (won't be great)
            logger.warning("PyMuPDF not available, PDF parsing limited")
            return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not available, DOCX parsing limited")
            return path.read_text(encoding="utf-8", errors="replace")

    # ── sandbox ───────────────────────────────────────────────────────

    def _check_sandbox(self, resolved: Path) -> None:
        """Ensure file is within allowed directories."""
        resolved_str = str(resolved)
        for allowed in self._config.allowed_dirs:
            allowed_resolved = str(Path(allowed).resolve())
            if resolved_str.startswith(allowed_resolved):
                return
        raise PermissionError(
            f"Access denied: {resolved} is outside allowed directories"
        )

    # ── format helpers ────────────────────────────────────────────────

    @staticmethod
    def _normalize_format(ext: str) -> str:
        mapping = {
            "txt": "txt",
            "text": "txt",
            "md": "md",
            "markdown": "md",
            "html": "html",
            "htm": "html",
            "pdf": "pdf",
            "docx": "docx",
        }
        fmt = mapping.get(ext)
        if fmt is None:
            raise ValueError(f"Unsupported file extension: .{ext}")
        return fmt

    # ── default summarizer ────────────────────────────────────────────

    @staticmethod
    def _default_summarizer(text: str) -> str:
        """Truncation-based fallback summary."""
        text = text.strip()
        if len(text) <= 500:
            return text
        return text[:500] + "..."
