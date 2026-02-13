"""DOCX document parser.

Prefers `python-docx` when available, but includes a lightweight fallback
implementation that extracts text directly from the DOCX ZIP container.
"""

import io
import zipfile
import xml.etree.ElementTree as ET

from bantz.document.parsers.base import DocumentParser, ParseResult


# DOCX magic bytes (ZIP file with specific structure)
DOCX_MAGIC = b"PK\x03\x04"


class DOCXParser(DocumentParser):
    """
    Parser for DOCX documents.
    
    Uses python-docx library for text extraction.
    """
    
    def __init__(self):
        """Initialize DOCX parser."""
        self._docx = None
        self._load_library()
    
    def _load_library(self) -> None:
        """Load python-docx library."""
        try:
            import docx
            self._docx = docx
        except ImportError:
            pass
    
    @property
    def is_available(self) -> bool:
        """DOCX parsing is always available (fallback is built-in)."""
        return True

    def _parse_with_fallback(self, data: bytes) -> tuple[str, dict]:
        """Extract text from DOCX bytes without external dependencies."""
        try:
            zf = zipfile.ZipFile(io.BytesIO(data))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX as zip: {e}") from e

        try:
            xml_bytes = zf.read("word/document.xml")
        except KeyError as e:
            raise ValueError("DOCX is missing word/document.xml") from e

        try:
            root = ET.fromstring(xml_bytes)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX XML: {e}") from e

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs: list[str] = []

        for p in root.findall(".//w:body//w:p", ns):
            parts: list[str] = []
            for t in p.findall(".//w:t", ns):
                if t.text:
                    parts.append(t.text)
            line = "".join(parts).strip()
            if line:
                paragraphs.append(line)

        return "\n\n".join(paragraphs), {}
    
    async def parse(self, data: bytes) -> ParseResult:
        """
        Parse DOCX document and extract text.
        
        Args:
            data: Raw DOCX bytes.
            
        Returns:
            ParseResult with extracted text.
            
        Raises:
            ValueError: If DOCX cannot be parsed.
        """
        if not self.can_parse(data):
            raise ValueError("Data does not appear to be a valid DOCX")
        
        text_parts = []
        metadata = {}
        
        try:
            if self._docx is not None:
                doc = self._docx.Document(io.BytesIO(data))

                # Extract core properties
                if doc.core_properties:
                    props = doc.core_properties
                    if props.author:
                        metadata["author"] = props.author
                    if props.title:
                        metadata["title"] = props.title
                    if props.created:
                        metadata["created"] = props.created.isoformat()
                    if props.modified:
                        metadata["modified"] = props.modified.isoformat()

                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
            else:
                fallback_text, fallback_meta = self._parse_with_fallback(data)
                if fallback_text.strip():
                    text_parts.append(fallback_text)
                metadata.update(fallback_meta)

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}") from e
        
        full_text = "\n\n".join(text_parts)
        
        # Estimate page count (rough: ~500 words per page)
        word_count = len(full_text.split())
        page_count = max(1, word_count // 500)
        
        return ParseResult(
            text=full_text,
            page_count=page_count,
            word_count=word_count,
            metadata=metadata,
        )
    
    def can_parse(self, data: bytes) -> bool:
        """
        Check if data is a DOCX.
        
        Args:
            data: Raw bytes to check.
            
        Returns:
            True if data appears to be a DOCX (ZIP with specific content).
        """
        if data[:4] != DOCX_MAGIC:
            return False
        
        # Further check: look for [Content_Types].xml or word/ in the ZIP
        # For performance, just check magic bytes
        return True
    
    @property
    def supported_extensions(self) -> list[str]:
        """Get supported file extensions."""
        return [".docx"]
    
    @property
    def mime_types(self) -> list[str]:
        """Get supported MIME types."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]


class DOCParser(DocumentParser):
    """
    Parser for legacy DOC documents.
    
    Note: Legacy DOC parsing is limited without external tools.
    This provides basic support only.
    """
    
    # DOC magic bytes (OLE Compound File)
    DOC_MAGIC = b"\xd0\xcf\x11\xe0"
    
    def __init__(self):
        """Initialize DOC parser."""
        self._antiword_available = False
        self._check_antiword()
    
    def _check_antiword(self) -> None:
        """Check if antiword is available for DOC parsing."""
        import shutil
        self._antiword_available = shutil.which("antiword") is not None
    
    @property
    def is_available(self) -> bool:
        """Check if DOC parsing is available."""
        return self._antiword_available
    
    async def parse(self, data: bytes) -> ParseResult:
        """
        Parse DOC document.
        
        Args:
            data: Raw DOC bytes.
            
        Returns:
            ParseResult with extracted text.
            
        Raises:
            ValueError: If DOC cannot be parsed.
        """
        if not self.is_available:
            raise ImportError(
                "antiword is not available. Install it: sudo apt install antiword"
            )
        
        if not self.can_parse(data):
            raise ValueError("Data does not appear to be a valid DOC")
        
        import subprocess
        import tempfile
        
        try:
            # Write to temp file and use antiword
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as f:
                f.write(data)
                f.flush()
                
                result = subprocess.run(
                    ["antiword", f.name],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                
                if result.returncode != 0:
                    raise ValueError(f"antiword failed: {result.stderr}")
                
                text = result.stdout
                
        except subprocess.TimeoutExpired:
            raise ValueError("DOC parsing timed out")
        except Exception as e:
            raise ValueError(f"Failed to parse DOC: {e}") from e
        
        return ParseResult(
            text=text,
            page_count=1,
        )
    
    def can_parse(self, data: bytes) -> bool:
        """Check if data is a DOC file."""
        return data[:4] == self.DOC_MAGIC
    
    @property
    def supported_extensions(self) -> list[str]:
        """Get supported file extensions."""
        return [".doc"]
    
    @property
    def mime_types(self) -> list[str]:
        """Get supported MIME types."""
        return ["application/msword"]


def create_docx_parser() -> DOCXParser:
    """Factory function to create a DOCX parser."""
    return DOCXParser()
